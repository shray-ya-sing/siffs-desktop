"""
Word metadata extractor using python-docx library.
Extracts comprehensive metadata about the document, including properties and content structure.
"""

import os
import json
from typing import Dict, Any, Optional
from datetime import datetime
from pathlib import Path
import logging

try:
    import docx
    from docx.document import Document as DocxDocument
    DOCX_AVAILABLE = True
    print("python-docx imported successfully")
except ImportError as e:
    DOCX_AVAILABLE = False
    print(f"python-docx import failed: {e}")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class WordMetadataExtractor:
    """
    Comprehensive Word metadata extractor using python-docx.
    Extracts document properties, structure, and content.
    """
    
    def __init__(self, document_path: Optional[str] = None):
        """
        Initialize the Word metadata extractor.
        
        Args:
            document_path: Path to the Word file (optional, can be set later)
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required but not installed. Please install with: pip install python-docx")
            
        self.document_path = document_path
        self.document: Optional[DocxDocument] = None
        
    def open_document(self, document_path: Optional[str] = None) -> None:
        """
        Open the Word document.
        
        Args:
            document_path: Path to the Word file
        """
        if document_path:
            self.document_path = document_path
            
        if not self.document_path:
            raise ValueError("No document path specified")
            
        if not os.path.exists(self.document_path):
            raise FileNotFoundError(f"Document not found: {self.document_path}")
            
        try:
            self.document = docx.Document(self.document_path)
        except Exception as e:
            raise Exception(f"Failed to open document: {str(e)}")
    
    def close(self) -> None:
        """Close the document and clean up resources."""
        self.document = None
    
    def __enter__(self):
        """Context manager entry."""
        return self
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        """Context manager exit - cleanup resources."""
        self.close()
        
    def extract_document_metadata(
        self, 
        document_path: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Extract comprehensive metadata from the Word document.
        
        Args:
            document_path: Path to the Word file
            
        Returns:
            Dictionary containing complete document metadata
        """
        if document_path or not self.document:
            self.open_document(document_path)
            
        try:
            # Get basic document info
            document_metadata = {
                "extractedAt": datetime.now().isoformat(),
                "documentPath": self.document_path,
                "documentName": os.path.basename(self.document_path) if self.document_path else "Unknown",
                "coreProperties": self._extract_core_properties(),
                "paragraphs": self._extract_paragraphs(),
                "tables": self._extract_tables()
            }
            
            return document_metadata
            
        except Exception as e:
            raise Exception(f"Failed to extract document metadata: {str(e)}")
    
    def _extract_core_properties(self) -> Dict[str, Any]:
        """Extract core document properties."""
        try:
            core_props = self.document.core_properties
            return {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "category": core_props.category or "",
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "lastModifiedBy": core_props.last_modified_by or "",
                "revision": core_props.revision,
                "version": core_props.version or ""
            }
        except Exception as e:
            logger.warning(f"Error extracting core properties: {str(e)}")
            return {"error": str(e)}
    
    def _extract_paragraphs(self) -> list:
        """Extract paragraphs and their text from the document."""
        try:
            return [paragraph.text for paragraph in self.document.paragraphs]
        except Exception as e:
            logger.warning(f"Error extracting paragraphs: {str(e)}")
            return [{"error": str(e)}]
    
    def _extract_tables(self) -> list:
        """Extract tables and their content from the document."""
        tables_data = []
        try:
            for table in self.document.tables:
                table_content = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_content.append(row_data)
                tables_data.append(table_content)
            return tables_data
        except Exception as e:
            logger.warning(f"Error extracting tables: {str(e)}")
            return [{"error": str(e)}]
    
    def extract_metadata(self, document_path: str) -> Dict[str, Any]:
        """
        Extract metadata for cache handler compatibility.
        
        Args:
            document_path: Path to the Word file
            
        Returns:
            Dictionary containing document metadata
        """
        return self.extract_document_metadata(document_path)
    
    def extract_to_json(
        self,
        document_path: Optional[str] = None,
        output_path: Optional[str] = None
    ) -> str:
        """
        Extract metadata and return as JSON string.
        
        Args:
            document_path: Path to the Word file
            output_path: Optional path to save JSON file
            
        Returns:
            JSON string containing the metadata
        """
        metadata = self.extract_document_metadata(document_path)
        json_str = json.dumps(metadata, indent=2, ensure_ascii=False)
        
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(json_str)
            logger.info(f"Metadata saved to: {output_path}")
            
        return json_str

